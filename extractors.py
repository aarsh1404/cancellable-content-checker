"""
Content extraction utilities for various file formats and URLs
Enhanced with visual content analysis capabilities
"""

import io
import requests
from typing import Optional, Union, Dict
from urllib.parse import urlparse
import streamlit as st

# Import file processing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

class ContentExtractor:
    """Utility class for extracting text content from various sources"""
    
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
        
        # Initialize visual analyzer
        try:
            from visual_analyzer import VisualURLAnalyzer
            self.visual_analyzer = VisualURLAnalyzer()
            self.visual_available = True
        except ImportError:
            self.visual_analyzer = None
            self.visual_available = False
    
    def extract_from_file(self, uploaded_file) -> str:
        """
        Extract text content from uploaded file
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Extracted text content
        """
        try:
            file_type = uploaded_file.type.lower()
            file_name = uploaded_file.name.lower()
            
            # Reset file pointer
            uploaded_file.seek(0)
            
            if file_type == 'text/plain' or file_name.endswith('.txt'):
                return self._extract_from_txt(uploaded_file)
            
            elif file_type == 'application/pdf' or file_name.endswith('.pdf'):
                return self._extract_from_pdf(uploaded_file)
            
            elif (file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' 
                  or file_name.endswith('.docx')):
                return self._extract_from_docx(uploaded_file)
            
            elif file_type in ['image/jpeg', 'image/jpg', 'image/png'] or file_name.endswith(('.jpg', '.jpeg', '.png')):
                return self._extract_from_image(uploaded_file)
            
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            raise Exception(f"Failed to extract content from file: {str(e)}")
    
    def _extract_from_txt(self, file) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    file.seek(0)
                    content = file.read().decode(encoding)
                    return content.strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            file.seek(0)
            content = file.read().decode('utf-8', errors='ignore')
            return content.strip()
            
        except Exception as e:
            raise Exception(f"Failed to read text file: {str(e)}")
    
    def _extract_from_pdf(self, file) -> str:
        """Extract text from PDF file"""
        if not PDF_AVAILABLE:
            raise Exception("PyPDF2 not available. Please install it with: pip install PyPDF2")
        
        try:
            file.seek(0)
            pdf_reader = PyPDF2.PdfReader(file)
            
            text_content = []
            for page_num in range(len(pdf_reader.pages)):
                try:
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
                except Exception as e:
                    st.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                    continue
            
            if not text_content:
                raise Exception("No readable text found in PDF")
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_from_docx(self, file) -> str:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not available. Please install it with: pip install python-docx")
        
        try:
            file.seek(0)
            doc = Document(file)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            if not text_content:
                raise Exception("No readable text found in DOCX")
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_from_image(self, file) -> str:
        """Extract text from image using OCR"""
        if not OCR_AVAILABLE:
            raise Exception("OCR not available. Please install: pip install Pillow pytesseract")
        
        try:
            file.seek(0)
            image = Image.open(file)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Extract text using OCR
            text = pytesseract.image_to_string(image)
            
            if not text.strip():
                raise Exception("No readable text found in image")
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Failed to extract text from image: {str(e)}")
    
    def extract_from_url(self, url: str) -> str:
        """
        Extract text content from URL
        
        Args:
            url: URL to extract content from
            
        Returns:
            Extracted text content
        """
        try:
            # Validate URL
            parsed_url = urlparse(url)
            if not parsed_url.scheme or not parsed_url.netloc:
                raise ValueError("Invalid URL format")
            
            # Make request with timeout
            response = self.session.get(url, timeout=10)
            response.raise_for_status()
            
            # Check if it's HTML content
            content_type = response.headers.get('content-type', '').lower()
            
            if 'text/html' in content_type:
                return self._extract_from_html(response.text, url)
            elif 'text/plain' in content_type:
                return response.text.strip()
            else:
                raise Exception(f"Unsupported content type: {content_type}")
                
        except requests.exceptions.RequestException as e:
            raise Exception(f"Failed to fetch URL: {str(e)}")
        except Exception as e:
            raise Exception(f"Failed to extract content from URL: {str(e)}")
    
    def extract_from_url_enhanced(self, url: str) -> Dict:
        """
        Enhanced URL extraction with visual content analysis
        
        Args:
            url: URL to analyze
            
        Returns:
            Dictionary with text content, visual elements, and metadata
        """
        if not self.visual_available or not self.visual_analyzer:
            # Fallback to basic extraction
            try:
                text_content = self.extract_from_url(url)
                return {
                    'text_content': text_content,
                    'visual_elements': [],
                    'screenshot': None,
                    'metadata': {},
                    'images': [],
                    'error': None
                }
            except Exception as e:
                return {
                    'text_content': '',
                    'visual_elements': [],
                    'screenshot': None,
                    'metadata': {},
                    'images': [],
                    'error': str(e)
                }
        
        # Use enhanced visual analyzer
        return self.visual_analyzer.analyze_url_enhanced(url)
    
    def _extract_from_html(self, html_content: str, url: str) -> str:
        """Extract text content from HTML"""
        if not BS4_AVAILABLE:
            raise Exception("BeautifulSoup not available. Please install it with: pip install beautifulsoup4")
        
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Try to find main content area
            main_content = None
            
            # Look for common main content selectors
            main_selectors = [
                'main', 'article', '.content', '#content', '.main-content',
                '.post-content', '.entry-content', '.article-content'
            ]
            
            for selector in main_selectors:
                element = soup.select_one(selector)
                if element:
                    main_content = element
                    break
            
            # If no main content found, use body
            if not main_content:
                main_content = soup.find('body') or soup
            
            # Extract text
            text = main_content.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            if not text.strip():
                raise Exception("No readable text found in HTML")
            
            # Limit text length for analysis
            if len(text) > 4000:
                text = text[:4000] + "..."
            
            return text
            
        except Exception as e:
            raise Exception(f"Failed to parse HTML: {str(e)}")
    
    def is_supported_file_type(self, file_type: str, file_name: str) -> bool:
        """Check if file type is supported"""
        supported_types = [
            'text/plain',
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'image/jpeg',
            'image/jpg',
            'image/png'
        ]
        
        supported_extensions = ['.txt', '.pdf', '.docx', '.jpg', '.jpeg', '.png']
        
        return (file_type.lower() in supported_types or 
                any(file_name.lower().endswith(ext) for ext in supported_extensions))
    
    def get_file_info(self, uploaded_file) -> dict:
        """Get information about uploaded file"""
        return {
            'name': uploaded_file.name,
            'size': uploaded_file.size,
            'type': uploaded_file.type,
            'supported': self.is_supported_file_type(uploaded_file.type, uploaded_file.name)
        }
    
    def validate_url(self, url: str) -> bool:
        """Validate URL format"""
        try:
            parsed = urlparse(url)
            return bool(parsed.scheme and parsed.netloc)
        except:
            return False
